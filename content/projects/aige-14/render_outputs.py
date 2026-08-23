#!/usr/bin/env python3
"""Render the reviewed Aige 14 DOCX to frozen Markdown and WeChat HTML."""

from __future__ import annotations

import argparse
import html
import re
from pathlib import Path

from docx import Document


SECTION_NUMBERS = set("一二三四五六七八九十")


def is_section_number(text: str) -> bool:
    return bool(text) and len(text) <= 2 and all(ch in SECTION_NUMBERS for ch in text)


def paragraph_is_bold(paragraph) -> bool:
    visible = [run for run in paragraph.runs if run.text.strip()]
    return bool(visible) and sum(bool(run.bold) for run in visible) >= len(visible) / 2


def leaf(text: str, style: str = "") -> str:
    style_attr = f' style="{style}"' if style else ""
    return f'<span leaf=""{style_attr}>{html.escape(text)}</span>'


def render_markdown(paragraphs, output: Path) -> None:
    lines = []
    index = 0
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        text = paragraph.text.strip()
        if not text:
            lines.extend(["<!-- wechat:blank -->", ""])
            index += 1
            continue
        if index == 0:
            lines.extend([text, ""])
        elif text == "让爱有方向，让生命专注":
            lines.extend([f"# {text}", ""])
        elif is_section_number(text):
            next_text = paragraphs[index + 1].text.strip() if index + 1 < len(paragraphs) else ""
            if next_text:
                lines.extend([f"## {text}　{next_text}", ""])
                index += 1
            else:
                lines.extend([f"## {text}", ""])
        elif paragraph_is_bold(paragraph):
            lines.extend([f"**{text}**", ""])
        else:
            lines.extend([text, ""])
        index += 1
    output.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def body_paragraph(text: str, bold: bool) -> str:
    weight = "font-weight:700;" if bold else "font-weight:400;"
    style = (
        "margin:0 0 16px 0;padding:0;color:#343832;font-size:16px;"
        f"line-height:1.9;letter-spacing:0.04em;{weight}text-align:justify;"
    )
    return f'<p style="{style}">{leaf(text)}</p>'


def render_html(paragraphs, output: Path) -> None:
    parts = [
        '<section style="max-width:677px;margin:0 auto;padding:22px 16px 46px 16px;'
        'background:#fbfbf7;color:#343832;font-family:-apple-system,BlinkMacSystemFont,'
        "'PingFang SC','Hiragino Sans GB','Microsoft YaHei',sans-serif;box-sizing:border-box;\">"
    ]

    eyebrow = paragraphs[0].text.strip()
    title = paragraphs[2].text.strip()
    subtitle = paragraphs[3].text.strip()
    parts.extend([
        '<section style="margin:0 0 34px 0;padding:22px 18px 24px 18px;'
        'border-left:4px solid #7b846d;background:#f1f2eb;box-sizing:border-box;">',
        f'<p style="margin:0 0 13px 0;color:#7b846d;font-size:12px;line-height:1.5;letter-spacing:0.18em;font-weight:600;">{leaf(eyebrow)}</p>',
        f'<h1 style="margin:0 0 10px 0;padding:0;color:#252a25;font-size:27px;line-height:1.35;letter-spacing:0.04em;font-weight:700;">{leaf(title)}</h1>',
        f'<p style="margin:0;color:#686e64;font-size:14px;line-height:1.75;letter-spacing:0.03em;">{leaf(subtitle)}</p>',
        '</section>',
    ])

    index = 4
    first_body = True
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        text = paragraph.text.strip()
        if not text:
            parts.append('<section style="height:9px;line-height:9px;"><br></section>')
            index += 1
            continue

        if is_section_number(text):
            next_text = paragraphs[index + 1].text.strip() if index + 1 < len(paragraphs) else ""
            parts.extend([
                '<section style="margin:38px 0 24px 0;padding:0 0 13px 0;border-bottom:1px solid #d8dacf;">',
                f'<p style="margin:0 0 7px 0;color:#89917d;font-size:12px;line-height:1;letter-spacing:0.18em;font-weight:600;">{leaf(text)}</p>',
                f'<h2 style="margin:0;padding:0;color:#2b302a;font-size:21px;line-height:1.5;letter-spacing:0.035em;font-weight:700;">{leaf(next_text)}</h2>',
                '</section>',
            ])
            index += 2
            first_body = False
            continue

        if first_body and index in {5, 6}:
            if index == 5:
                parts.append('<section style="margin:0 0 27px 0;padding:20px 18px 4px 18px;background:#f3f4ee;border-radius:3px;">')
            parts.append(body_paragraph(text, paragraph_is_bold(paragraph)))
            if index == 6:
                parts.append('</section>')
            index += 1
            continue

        parts.append(body_paragraph(text, paragraph_is_bold(paragraph)))
        index += 1

    parts.extend([
        '<!-- generated:start -->',
        '<section style="margin:40px 0 0 0;text-align:center;">',
        f'<p style="margin:0;color:#9aa08f;font-size:11px;line-height:1.4;letter-spacing:0.28em;">{leaf("END")}</p>',
        '</section>',
        '<!-- generated:end -->',
        '</section>',
    ])
    output.write_text("\n".join(parts) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("markdown", type=Path)
    parser.add_argument("html", type=Path)
    args = parser.parse_args()

    doc = Document(args.source)
    args.markdown.parent.mkdir(parents=True, exist_ok=True)
    args.html.parent.mkdir(parents=True, exist_ok=True)
    render_markdown(doc.paragraphs, args.markdown)
    render_html(doc.paragraphs, args.html)
    print(f"wrote: {args.markdown}")
    print(f"wrote: {args.html}")


if __name__ == "__main__":
    main()
