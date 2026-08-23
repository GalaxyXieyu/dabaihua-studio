#!/usr/bin/env python3
"""Process Aige 14 through punctuation review and mobile-layout cleanup."""

from __future__ import annotations

import argparse
import copy
import difflib
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PUNCTUATION_EDITS = {
    "戴上之后，他成为孙悟空。虽说有了法力。但也有了约束。":
        "戴上之后，他成为孙悟空。虽说有了法力，但也有了约束。",
    "换个说法，他有了生的责任。也有了值得舍生取义的使命。":
        "换个说法，他有了生的责任，也有了值得舍生取义的使命。",
    "选择责任之后真正发生变化的，是他的力量，开始不再四处散落。":
        "选择责任之后真正发生变化的，是他的力量开始不再四处散落。",
    "而是：减少耗散，让人专注。":
        "而是减少耗散，让人专注。",
    "而是：我们没有把力量真正交给那些重要的东西。":
        "而是我们没有把力量真正交给那些重要的东西。",
    "“逝者如斯不舍昼夜”。":
        "“逝者如斯，不舍昼夜。”",
    "不仅可以理解时间。":
        "不仅可以理解时间，",
    "只有形成河岸，水流才会变深。才会形成方向。才可能奔腾。":
        "只有形成河岸，水流才会变深，才会形成方向，才可能奔腾。",
    "真正的立志不是因为：“现在的我还不够好，所以我要成为另外一个人。”":
        "真正的立志不是因为“现在的我还不够好，所以我要成为另外一个人”。",
    "而更像是：“我已经开始看见自己是谁，因此愿意忠实地把这个生命活出来。”":
        "而更像是“我已经开始看见自己是谁，因此愿意忠实地把这个生命活出来”。",
    "一个真正有使命的人，好像应该永远奔波。应该不断牺牲。应该把所有个人生活置于更宏大的目标之后。":
        "一个真正有使命的人，好像应该永远奔波，应该不断牺牲，应该把所有个人生活置于更宏大的目标之后。",
    "真正的爱，不只是：我愿意为你做什么。":
        "真正的爱，不只是我愿意为你做什么。",
    "而是能够说：我愿意陪伴你。愿意支持你。":
        "而是能够说：我愿意陪伴你，愿意支持你。",
    "真正的知，最后一定会进入行动。进入时间。进入关系。进入身体。":
        "真正的知，最后一定会进入行动，进入时间，进入关系，进入身体。",
    "然后是爱。我要寻找，也学习成为那个：能够陪伴别人、也让别人陪伴自己成为自己的人。":
        "然后是爱。我要寻找，也学习成为那个能够陪伴别人、也让别人陪伴自己成为自己的人。",
    "再后来，是苏东坡的风雨，和奥德赛的归途：即使身份、成败、环境都发生变化：我还能不能安住自己？":
        "再后来，是苏东坡的风雨，和奥德赛的归途：即使身份、成败、环境都发生变化，我还能不能安住自己？",
    "当我开始知道自己真正爱什么：我愿意把这份爱带向哪里？":
        "当我开始知道自己真正爱什么，我愿意把这份爱带向哪里？",
    "既然已经知道什么值得：我愿意把有限的力量真正交给什么？":
        "既然已经知道什么值得，我愿意把有限的力量真正交给什么？",
    "真正重要的是：不要浪费已经拥有的力量。":
        "真正重要的是不要浪费已经拥有的力量。",
    "真正重要的是：让真正的爱进入真实而持续的行动。":
        "真正重要的是让真正的爱进入真实而持续的行动。",
}


def copy_rpr(run):
    return copy.deepcopy(run._element.rPr) if run._element.rPr is not None else None


def source_chars(paragraph):
    chars = []
    for run in paragraph.runs:
        props = copy_rpr(run)
        for char in run.text:
            chars.append((char, props))
    return chars


def remove_runs(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)


def normalize_fonts(doc):
    """Replace unavailable legacy CJK fonts with a stable macOS CJK font."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Songti SC"
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), "Songti SC")
            rfonts.set(qn("w:hAnsi"), "Songti SC")
            rfonts.set(qn("w:eastAsia"), "Songti SC")
            rfonts.set(qn("w:cs"), "Songti SC")


def props_key(props) -> bytes:
    return b"" if props is None else props.xml.encode("utf-8")


def rebuild_paragraph(paragraph, new_text: str):
    old = paragraph.text
    chars = source_chars(paragraph)
    matcher = difflib.SequenceMatcher(a=old, b=new_text, autojunk=False)
    rebuilt = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            rebuilt.extend((chars[i][0], chars[i][1], False) for i in range(i1, i2))
            continue

        # The character immediately before an edited mark is blue for review.
        if rebuilt:
            previous = rebuilt[-1]
            rebuilt[-1] = (previous[0], previous[1], True)

        fallback = None
        if i1 < len(chars):
            fallback = chars[i1][1]
        elif chars:
            fallback = chars[-1][1]

        for char in new_text[j1:j2]:
            rebuilt.append((char, fallback, True))

    remove_runs(paragraph)
    current_run = None
    current_key = None
    for char, props, blue in rebuilt:
        key = (props_key(props), blue)
        if current_run is None or key != current_key:
            current_run = paragraph.add_run()
            if props is not None:
                if current_run._element.rPr is not None:
                    current_run._element.remove(current_run._element.rPr)
                current_run._element.insert(0, copy.deepcopy(props))
            if blue:
                current_run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            current_key = key
        current_run.add_text(char)


def apply_punctuation(source: Path, output: Path):
    doc = Document(source)
    found = set()
    for paragraph in doc.paragraphs:
        original = paragraph.text
        replacement = PUNCTUATION_EDITS.get(original)
        if replacement is None:
            continue
        rebuild_paragraph(paragraph, replacement)
        found.add(original)

    if len(found) != len(PUNCTUATION_EDITS):
        missing = sorted(set(PUNCTUATION_EDITS) - found)
        raise RuntimeError(f"Missing expected paragraphs: {missing}")

    normalize_fonts(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def is_section_number(text: str) -> bool:
    return text in "一二三四五六七八九十" and len(text) <= 2


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def apply_mobile_layout(source: Path, output: Path):
    doc = Document(source)

    # Preserve intentional chapter breathing room, but collapse accidental runs
    # of more than two empty paragraphs and remove trailing empty paragraphs.
    empty_run = 0
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            empty_run = 0
            continue
        empty_run += 1
        if empty_run > 2:
            remove_paragraph(paragraph)

    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        remove_paragraph(doc.paragraphs[-1])

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        fmt = paragraph.paragraph_format
        if not text:
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.line_spacing = 1
            continue

        if index <= 3:
            fmt.keep_with_next = index < 3
            continue

        if is_section_number(text):
            fmt.space_before = Pt(18)
            fmt.space_after = Pt(4)
            fmt.keep_with_next = True
            continue

        previous = doc.paragraphs[index - 1].text.strip() if index else ""
        if is_section_number(previous):
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(12)
            fmt.keep_with_next = True
            continue

        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = 1.65
        fmt.space_after = Pt(7)

    normalize_fonts(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("punctuation", type=Path)
    parser.add_argument("layout", type=Path)
    args = parser.parse_args()
    apply_punctuation(args.source, args.punctuation)
    apply_mobile_layout(args.punctuation, args.layout)
    print(f"punctuation edits: {len(PUNCTUATION_EDITS)}")
    print(f"wrote: {args.punctuation}")
    print(f"wrote: {args.layout}")


if __name__ == "__main__":
    main()
