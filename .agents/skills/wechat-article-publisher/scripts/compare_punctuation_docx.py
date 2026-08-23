#!/usr/bin/env python3
"""Compare two DOCX files and report paragraph-level punctuation changes."""

from __future__ import annotations

import argparse
import json
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document


PUNCTUATION = set("，。；：、？！“”‘’《》（）〔〕【】———…·,.!?;:'\"()[]{}-–—")


def blue_runs(paragraph):
    found = []
    for index, run in enumerate(paragraph.runs):
        color = run.font.color.rgb if run.font.color else None
        if color and str(color).upper() in {"0000FF", "0070C0", "0563C1", "2F5496"}:
            found.append({"run": index, "text": run.text, "color": str(color)})
    return found


def punctuation_only(before: str, after: str) -> bool:
    strip = lambda text: "".join(char for char in text if char not in PUNCTUATION)
    return strip(before) == strip(after)


def compare(original: Path, edited: Path):
    left = Document(original)
    right = Document(edited)
    left_text = [p.text for p in left.paragraphs]
    right_text = [p.text for p in right.paragraphs]
    matcher = SequenceMatcher(None, left_text, right_text, autojunk=False)
    changes = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        before = left_text[i1:i2]
        after = right_text[j1:j2]
        changes.append(
            {
                "type": tag,
                "original_range": [i1, i2],
                "edited_range": [j1, j2],
                "original": before,
                "edited": after,
                "punctuation_only": punctuation_only("".join(before), "".join(after)),
                "blue_runs": [blue_runs(right.paragraphs[j]) for j in range(j1, j2)],
            }
        )
    return {
        "original": str(original),
        "edited": str(edited),
        "original_paragraphs": len(left_text),
        "edited_paragraphs": len(right_text),
        "changes": changes,
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("original", type=Path)
    parser.add_argument("edited", type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    report = compare(args.original, args.edited)
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
        return
    print(
        f"原稿段落：{report['original_paragraphs']}；"
        f"校订稿段落：{report['edited_paragraphs']}；"
        f"差异组：{len(report['changes'])}"
    )
    for change in report["changes"]:
        print(
            f"\n[{change['type']}] "
            f"原稿{change['original_range']} -> 校订稿{change['edited_range']} "
            f"仅标点={change['punctuation_only']}"
        )
        for text in change["original"]:
            print(f"- {text}")
        for text in change["edited"]:
            print(f"+ {text}")


if __name__ == "__main__":
    main()
